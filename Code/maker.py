import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
#key
#%=newLine
#{=start var
#}=end var
#"#"number-=identify and start section
#~=end section
#[=start comment
#]=end comment
#\=nullify special char

class Maker:
    #input files
    inpF=[]
    #the document as a list
    docL=[]
    #path for the folder that contains files
    path=""
    #special characters
    key={"newLine":'%',
        "startVar":'{',
        "endVar":'}',
        "identSec":'#',
        "startSec":'-',
        "startComm":'[',
        "endComm":']',
        "nullify":'\\'}
    #store varaibles
    vars={}
    #jobFileContents
    jobFile={}

    #constructor, takes path for the folder that contains files
    def __init__(self, tpath):
        self.path=tpath

    #reads in the input files and generates a Doc
    def run(self):
        #prompt user to use job file
        inp=input("Take Inputs from file?\ny or n\n")

        #if yes, ask for file and read it in
        print("Running DocMaker...")
        if(inp=='y'):
            self.readJobF()
            for i in self.inpF:
                print("Opening File "+i+"...")
                self.readF(i)
        #if not get input files
        else:
            inp=input("Enter name of file to read from OR \"end\" to finish: ")
            while(inp!="end"):
                print("Opening File "+inp+"...")
                self.readF(inp)
                inp=input("Enter name of file to read from OR \"end\" to finish: ")

        #prompt user for output type
        docType=self.takeInp("Print to Microsoft Word Document?\ny or n\n", "Word?")
        dname=self.takeInp("Enter Name Of Document: ", "Doc Name")
        title=self.takeInp("Enter a Tile, or 'n' if you dont want one: ", "Title")
        if(docType=='y'):
            self.writeLetterWord(title, dname)
        else:
            self.writeLetterText(title, dname)
        print("DocMaker Complete!")

    #transfers docL from array to Microsoft Word Doc
    def writeLetterWord(self, title, dname):
        myDoc=docx.Document()
        if(title!='n'):
            titlePar=myDoc.add_paragraph(title)
            titlePar.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.CENTER
        for i in self.docL:
            par=myDoc.add_paragraph(i)
        myDoc.save(self.path+"../Docs/"+dname+".docx")
        print("File, "+ dname+" Doc.docx, Complete!")

    #transfers docL to .txt file
    def writeLetterText(self, title, dname):
        doc=""
        if(title!='n'):
            doc=title
        myFile=open(self.path+dname+" Doc.txt","w")
        for i in self.docL:
            doc=doc+i+"\n"
        myFile.write(doc)
        print("File, "+ dname+" Doc.docx, Complete!")
        myFile.close()
    
    #reads a file, and adds the correct contents to docL
    def readF(self, fname):
        fCont=""
        f=open(self.path+fname)
        for line in f:

            #removes extraneous newlines and adds to file
            if(line[-1]=="\n"):
                line=line[:-1]
            fCont=fCont+line
        f.close()

        #identifies if it has sections, and if so call choose
        if(fCont[0]=='#'):
            fCont=self.choose(fCont, fname)

        #calls parse to fill varaibles
        self.docL.append(self.parse(fCont))

    def readJobF(self):
        f=open(self.path+input("Enter File Name: "))
        for line in f:
            #removes extraneous newlines and adds to file
            if(line[-1]=="\n"):
                line=line[:-1]

            #parses the files line
            if('[' in line):
                line=line[:-1]
                line=line[1:]
                print(line)
                self.inpF=line.split(',')
            #seperates the variables name and value
            varName=""
            varVal=""
            nameEnded=False
            leadingSpace=False
            for i in line:
                if(i==':'):
                    nameEnded=True
                    leadingSpace=True
                elif(nameEnded):
                    if(leadingSpace and i==' '):
                        continue
                    leadingSpace=False
                    varVal=varVal+i
                else:
                    varName=varName+i
            self.jobFile[varName]=varVal
        f.close()

    #prompts the user for the correct section of the file, then returns only that section
    def choose(self, f, fname):
        #prompt user for section
        secNum=self.takeInp("Choose Section from " + fname + ": ", fname)
        lf=f.split('#')
        ind=0
        #go thorough array
        for i in lf:
            if(i==""):
                ind=ind+1
                continue

            #if there was a \# combine the sections
            if(i[-1]=='\\' and len(lf)>ind+1):
                lf[ind]=lf[ind]+'#'+lf[ind+1]

            #test if section is correct one and return if so
            same=True
            for j in range(len(secNum)):
                if(secNum[j]!=i[j]):
                    same=False
            if(same):
                return lf[ind][len(secNum)+1:]
            ind=ind+1
        raise Exception("Failed to Find Section "+secNum)



    #finds special characters and varaibles, and replaces them with the correct output, prompting the user if neccesary
    def parse(self, f):
        ind=0
        out=""
        offset=0
        for cha in f:
            #if no offset then add to output
            if(offset==0):

                #identifies nullified speical character and adds them unchanged, and idents tabs
                if(cha==self.key["nullify"]):
                    if(f[ind+1] in self.key.values()):
                        out=out+f[ind+1]
                    elif(f[ind+1]=='t'):
                        out=out+'\t'
                    offset=offset+1

                #identifies "%" and adds newline to the output
                elif(cha==self.key["newLine"]):
                    out=out+"\n"

                #finds the start of varaibles and prompts user for what to fill in, then adds them
                #to the output, and puts the length of the var in offset
                elif(cha==self.key["startVar"]):
                    offset=1
                    var=""
                    while(f[ind+offset]!=self.key["endVar"]):
                        var=var+f[ind+offset]
                        offset=offset+1
                    #add var to vars dict and prompt for value, or get value from vars dict
                    if(not(var in self.vars)):
                        inp=self.takeInp("Enter "+var+": ", var)
                        self.vars[var]=inp

                    out=out+self.vars[var]

                #ignore comments, by counting their length in offset
                elif(cha==self.key["startComm"]):
                    offset=1
                    var=""
                    while(f[ind+offset]!=self.key["endComm"]):
                        var=var+f[ind+offset]
                        offset=offset+1
                else:
                    out=out+cha
            else:
                offset=offset-1
                
            ind=ind+1

        #replace "a" with "an" if a vowel follows it
        vowels=["a","e","i","o","u","A","E","I","O","U"]
        for i in vowels:
            out=re.sub(" a "+i, " an "+i, out)
            out=re.sub("^A "+i, "An "+i, out)
            out=re.sub("A "+i, "An "+i, out)

        return out



    #replaces a section of a string
    def replace(self,string,sInd,fInd,newSec):
        start=string[:sInd]
        end=string[fInd:]
        return start+newSec+end

    #decides whether to take an input fomr command line or file
    def takeInp(self, message, key):
        if(key in self.jobFile):
            out=self.jobFile[key]
            print(message+out)
            return out
        else:
            return input(message)
